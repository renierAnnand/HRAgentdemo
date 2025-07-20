import streamlit as st
import fitz  # PyMuPDF for PDF reading
from docx import Document
import re
import time
import datetime
import json
import random

# Page configuration
st.set_page_config(
    page_title="AI Agent Demo - ARIA HR Assistant", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state for AI Agent conversation
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'agent_step' not in st.session_state:
    st.session_state.agent_step = 0
if 'extracted_data' not in st.session_state:
    st.session_state.extracted_data = None
if 'validation_complete' not in st.session_state:
    st.session_state.validation_complete = False

# Custom CSS for AI Agent interface
st.markdown("""
<style>
    .agent-message {
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border-left: 4px solid #4CAF50;
    }
    .agent-thinking {
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        background-color: #f8f9fa;
        border-left: 4px solid #ffc107;
        font-style: italic;
    }
    .agent-analysis {
        padding: 1rem;
        border-radius: 10px;
        margin: 0.5rem 0;
        background-color: #e8f4fd;
        border-left: 4px solid #2196F3;
    }
    .confidence-high { color: #4CAF50; font-weight: bold; }
    .confidence-medium { color: #FF9800; font-weight: bold; }
    .confidence-low { color: #f44336; font-weight: bold; }
    
    .chat-container {
        max-height: 400px;
        overflow-y: auto;
        padding: 1rem;
        border: 1px solid #ddd;
        border-radius: 10px;
        background-color: #fafafa;
    }
</style>
""", unsafe_allow_html=True)

# Sidebar - AI Agent Profile
with st.sidebar:
    st.markdown("# 🤖 ARIA")
    st.markdown("### AI Recruitment Intelligence Assistant")
    
    st.markdown("""
    **Status:** 🟢 Online  
    **Model:** GPT-HR-4 Turbo  
    **Specialization:** HR Document Processing  
    **Confidence Level:** 94.2%  
    **Documents Processed:** 15,847  
    """)
    
    st.markdown("---")
    st.markdown("### 🧠 Current Capabilities")
    capabilities = [
        "✅ Multi-format document reading",
        "✅ Intelligent field extraction", 
        "✅ Contextual data validation",
        "✅ ERP system integration",
        "✅ Compliance checking",
        "✅ Anomaly detection",
        "✅ Smart recommendations"
    ]
    for cap in capabilities:
        st.markdown(cap)
    
    st.markdown("---")
    st.markdown("### 📊 Session Analytics")
    st.metric("Processing Speed", "1.2s", "↓ 0.3s")
    st.metric("Accuracy Rate", "97.8%", "↑ 2.1%")
    st.metric("Fields Detected", "8/8", "↑ 2")

# Header
col1, col2 = st.columns([3, 1])
with col1:
    st.title("🤖 ARIA - AI HR Processing Agent")
    st.markdown("*Your intelligent assistant for automated new hire document processing*")

with col2:
    st.markdown("### 🔄 Agent Status")
    status_placeholder = st.empty()
    status_placeholder.success("🟢 Ready to Assist")

# Main interface
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown("## 📄 Document Upload")
    uploaded_file = st.file_uploader(
        "Upload New Hire Document", 
        type=["pdf", "docx"],
        help="I can process PDF and Word documents containing new hire information"
    )

with col2:
    st.markdown("## 💬 AI Agent Chat")
    chat_container = st.container()

def agent_message(message, message_type="info"):
    """Add an AI agent message to the chat"""
    timestamp = datetime.datetime.now().strftime("%H:%M:%S")
    
    if message_type == "thinking":
        css_class = "agent-thinking"
        icon = "🤔"
    elif message_type == "analysis":
        css_class = "agent-analysis" 
        icon = "🔍"
    else:
        css_class = "agent-message"
        icon = "🤖"
    
    st.session_state.messages.append({
        "message": message,
        "type": message_type,
        "timestamp": timestamp,
        "icon": icon
    })

def display_chat():
    """Display the AI agent chat interface"""
    with chat_container:
        st.markdown('<div class="chat-container">', unsafe_allow_html=True)
        
        for msg in st.session_state.messages[-10:]:  # Show last 10 messages
            if msg["type"] == "thinking":
                st.markdown(f'<div class="agent-thinking">{msg["icon"]} **[{msg["timestamp"]}]** {msg["message"]}</div>', unsafe_allow_html=True)
            elif msg["type"] == "analysis":
                st.markdown(f'<div class="agent-analysis">{msg["icon"]} **[{msg["timestamp"]}]** {msg["message"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="agent-message">{msg["icon"]} **[{msg["timestamp"]}]** {msg["message"]}</div>', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)

# AI Agent Functions
def ai_extract_text(file, filename):
    """AI-powered text extraction with reasoning"""
    agent_message("📄 Document received. Analyzing file structure...", "thinking")
    time.sleep(1)
    
    try:
        if filename.endswith(".pdf"):
            agent_message("🔍 Detected PDF format. Using advanced OCR analysis...", "analysis")
            text = ""
            doc = fitz.open(stream=file.read(), filetype="pdf")
            pages = len(doc)
            agent_message(f"📊 Document analysis: {pages} pages detected. Processing each page...", "analysis")
            
            for i, page in enumerate(doc):
                page_text = page.get_text()
                text += page_text
                if i == 0:  # Analyze first page structure
                    agent_message(f"🧠 Page 1 analysis: Found {len(page_text.split())} words, detecting form structure...", "thinking")
            
            doc.close()
            agent_message(f"✅ PDF processing complete. Extracted {len(text.split())} total words.", "info")
            
        else:
            agent_message("📝 Detected Word document. Parsing document structure...", "analysis")
            text = ""
            doc = Document(file)
            para_count = len(doc.paragraphs)
            agent_message(f"📊 Document structure: {para_count} paragraphs identified.", "analysis")
            
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            agent_message(f"✅ Word document processing complete. Content successfully extracted.", "info")
        
        return text
        
    except Exception as e:
        agent_message(f"❌ Error during document processing: {str(e)}", "info")
        return ""

def ai_analyze_and_extract(text):
    """AI-powered intelligent field extraction with reasoning"""
    agent_message("🧠 Starting intelligent field analysis...", "thinking")
    time.sleep(1)
    
    # AI analyzes document structure first
    lines = text.split('\n')
    non_empty_lines = [line.strip() for line in lines if line.strip()]
    
    agent_message(f"📊 Document structure analysis: {len(non_empty_lines)} content lines identified.", "analysis")
    
    # Check for form-like structure
    colon_count = sum(1 for line in non_empty_lines if ':' in line)
    agent_message(f"🔍 Detected structured format: {colon_count} field-value pairs found.", "thinking")
    
    fields = {}
    confidence_scores = {}
    
    # Enhanced AI patterns with confidence scoring
    field_patterns = {
        "Name": {
            "patterns": [
                r"(?:Name|Full Name|Employee Name):\s*([^\n\r]+)",
                r"(?:Name|Full Name|Employee Name)\s*[:\-]\s*([^\n\r]+)",
                r"Name\s+([A-Za-z\s]{3,})(?:\n|$)"
            ],
            "priority": 1
        },
        "Email": {
            "patterns": [
                r"(?:Email|Email Address|E-mail):\s*([^\s\n\r]+@[^\s\n\r]+)",
                r"([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})"
            ],
            "priority": 1
        },
        "Department": {
            "patterns": [
                r"(?:Department|Dept):\s*([^\n\r]+)",
                r"(?:Department|Dept)\s*[:\-]\s*([^\n\r]+)"
            ],
            "priority": 2
        },
        "Role": {
            "patterns": [
                r"(?:Role|Position|Job Title|Title):\s*([^\n\r]+)",
                r"(?:Role|Position|Job Title|Title)\s*[:\-]\s*([^\n\r]+)"
            ],
            "priority": 2
        },
        "Start Date": {
            "patterns": [
                r"(?:Start Date|Start|Begin Date|Commencement Date):\s*([^\n\r]+)",
                r"(?:Start Date|Start|Begin Date|Commencement Date)\s*[:\-]\s*([^\n\r]+)"
            ],
            "priority": 1
        },
        "Salary": {
            "patterns": [
                r"(?:Salary|Annual Salary|Compensation):\s*([^\n\r]+)",
                r"\$([0-9,]+(?:\.[0-9]{2})?)"
            ],
            "priority": 3
        },
        "Manager": {
            "patterns": [
                r"(?:Manager|Supervisor|Reports To):\s*([^\n\r]+)",
                r"(?:Manager|Supervisor|Reports To)\s*[:\-]\s*([^\n\r]+)"
            ],
            "priority": 3
        },
        "Employee ID": {
            "patterns": [
                r"(?:Employee ID|EMP ID|ID|Employee Number):\s*([^\n\r]+)",
                r"EMP\d+",
                r"(?:ID|Employee)\s*[:\-]\s*([A-Z0-9]+)"
            ],
            "priority": 3
        }
    }
    
    agent_message("🔍 Applying neural pattern recognition to extract fields...", "analysis")
    
    for field_name, field_info in field_patterns.items():
        best_match = None
        best_confidence = 0
        
        for i, pattern in enumerate(field_info["patterns"]):
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = match.group(1).strip() if len(match.groups()) > 0 else match.group(0).strip()
                if value and len(value) > 0:
                    # Calculate confidence based on pattern specificity and context
                    confidence = 0.9 - (i * 0.1)  # First patterns are more specific
                    if ':' in match.group(0):  # Structured format bonus
                        confidence += 0.1
                    
                    if confidence > best_confidence:
                        best_match = value
                        best_confidence = confidence
        
        if best_match:
            fields[field_name] = best_match
            confidence_scores[field_name] = best_confidence
            conf_color = "high" if best_confidence > 0.8 else "medium" if best_confidence > 0.6 else "low"
            agent_message(f"✅ {field_name}: <span class='confidence-{conf_color}'>Found with {best_confidence:.1%} confidence</span>", "analysis")
        else:
            fields[field_name] = "❌ Not Found"
            confidence_scores[field_name] = 0
            agent_message(f"⚠️ {field_name}: Not detected in document", "thinking")
    
    # AI reasoning about data quality
    found_fields = sum(1 for v in fields.values() if "Not Found" not in v)
    total_fields = len(fields)
    
    agent_message(f"🧠 Extraction complete: {found_fields}/{total_fields} fields successfully identified.", "info")
    
    if found_fields >= 5:
        agent_message("✅ High-quality data extraction achieved. Proceeding to validation phase.", "info")
    elif found_fields >= 3:
        agent_message("⚠️ Moderate data extraction. Some manual verification may be needed.", "info")
    else:
        agent_message("❌ Low extraction success. Document may need reformatting.", "info")
    
    return fields, confidence_scores

def ai_validate_data(fields, confidence_scores):
    """AI-powered data validation with intelligent reasoning"""
    agent_message("🧠 Initiating intelligent data validation sequence...", "thinking")
    time.sleep(1)
    
    validation_results = {}
    errors = []
    warnings = []
    suggestions = []
    
    # AI validation logic
    required_fields = ["Name", "Email", "Department", "Role", "Start Date"]
    
    agent_message("🔍 Validating required fields and data integrity...", "analysis")
    
    for field in required_fields:
        if "Not Found" in fields.get(field, "") or not fields.get(field, "").strip():
            errors.append(f"Missing critical field: {field}")
            validation_results[field] = "❌ Missing"
            agent_message(f"❌ Critical validation failure: {field} is required but missing", "analysis")
        else:
            validation_results[field] = "✅ Valid"
    
    # Advanced AI validation
    if fields.get("Email") and "Not Found" not in fields["Email"]:
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if re.match(email_pattern, fields["Email"]):
            agent_message(f"✅ Email validation passed: {fields['Email']} is properly formatted", "analysis")
            
            # AI domain analysis
            domain = fields["Email"].split('@')[1]
            if domain in ['gmail.com', 'yahoo.com', 'hotmail.com']:
                warnings.append("Personal email domain detected - consider using corporate email")
                suggestions.append("Request corporate email address for official records")
        else:
            errors.append("Email format validation failed")
            validation_results["Email"] = "❌ Invalid format"
            agent_message("❌ Email format validation failed - invalid structure detected", "analysis")
    
    # AI-powered date intelligence
    if fields.get("Start Date") and "Not Found" not in fields["Start Date"]:
        agent_message("🤖 Analyzing start date with temporal intelligence...", "thinking")
        date_str = fields["Start Date"]
        parsed_date = None
        
        for fmt in ["%Y-%m-%d", "%m/%d/%Y", "%d/%m/%Y", "%B %d, %Y", "%m-%d-%Y"]:
            try:
                parsed_date = datetime.datetime.strptime(date_str, fmt)
                break
            except ValueError:
                continue
        
        if parsed_date:
            days_from_now = (parsed_date.date() - datetime.date.today()).days
            
            if days_from_now < 0:
                warnings.append(f"Start date is {abs(days_from_now)} days in the past")
                agent_message(f"⚠️ Temporal anomaly: Start date is {abs(days_from_now)} days ago", "analysis")
            elif days_from_now > 90:
                warnings.append(f"Start date is {days_from_now} days in the future")
                suggestions.append("Verify if this is a future hire or if date needs correction")
            else:
                agent_message(f"✅ Start date validation passed: {days_from_now} days from today", "analysis")
            
            validation_results["Start Date"] = "✅ Valid"
        else:
            warnings.append("Could not parse start date format")
            validation_results["Start Date"] = "⚠️ Format unclear"
            suggestions.append("Standardize date format to YYYY-MM-DD")
    
    # AI salary analysis
    if fields.get("Salary") and "Not Found" not in fields["Salary"]:
        agent_message("💰 Running salary analysis with market intelligence...", "thinking")
        salary_str = re.sub(r'[,$]', '', fields["Salary"])
        try:
            salary_num = float(salary_str)
            
            # AI market analysis
            if salary_num < 25000:
                warnings.append("Salary below market minimum - verify accuracy")
                agent_message(f"⚠️ Salary alert: ${salary_num:,.0f} is below typical market rates", "analysis")
            elif salary_num > 300000:
                warnings.append("Executive-level salary detected - additional approvals may be needed")
                suggestions.append("Route through executive compensation review")
                agent_message(f"💼 Executive compensation detected: ${salary_num:,.0f}", "analysis")
            else:
                agent_message(f"✅ Salary validation passed: ${salary_num:,.0f} within normal range", "analysis")
            
            validation_results["Salary"] = "✅ Valid"
        except ValueError:
            warnings.append("Could not parse salary amount")
            validation_results["Salary"] = "⚠️ Invalid format"
    
    # AI reasoning summary
    total_issues = len(errors) + len(warnings)
    if total_issues == 0:
        agent_message("🎉 Validation complete: All data passes AI quality checks!", "info")
    else:
        agent_message(f"📊 Validation summary: {len(errors)} errors, {len(warnings)} warnings identified", "info")
    
    if suggestions:
        agent_message(f"💡 AI generated {len(suggestions)} optimization suggestions", "info")
    
    return validation_results, errors, warnings, suggestions

def ai_erp_integration(fields):
    """AI-powered ERP integration simulation"""
    agent_message("🚀 Initiating AI-driven ERP integration sequence...", "info")
    time.sleep(1)
    
    # AI system selection
    erp_systems = ["Oracle HCM Cloud", "SAP SuccessFactors", "Workday HCM"]
    selected_erp = random.choice(erp_systems)
    agent_message(f"🤖 AI selected optimal ERP system: {selected_erp}", "analysis")
    
    integration_steps = [
        ("🔐 Establishing secure API connection", 2),
        ("🧠 AI mapping employee data to ERP schema", 2),
        ("🔍 Running pre-integration data validation", 1),
        ("⚡ Creating employee profile with AI optimization", 2),
        ("🛡️ Configuring security roles and permissions", 1),
        ("💰 Setting up payroll integration", 2),
        ("📧 Triggering automated welcome workflow", 1),
        ("📊 Updating organizational charts and reporting", 1),
        ("🔄 Syncing with downstream systems", 1)
    ]
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    for i, (step, duration) in enumerate(integration_steps):
        agent_message(step, "thinking")
        status_text.text(f"AI Agent: {step}")
        time.sleep(duration)
        progress_bar.progress((i + 1) / len(integration_steps))
    
    status_text.empty()
    progress_bar.empty()
    
    # AI generates insights
    emp_id = fields.get("Employee ID", f"AI{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}")
    
    agent_message(f"✅ Integration complete! Employee {emp_id} successfully onboarded.", "info")
    agent_message("🧠 AI has triggered 7 downstream automation workflows", "analysis")
    
    return {
        "employee_id": emp_id,
        "status": "SUCCESS",
        "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "system": selected_erp,
        "ai_insights": [
            "Optimal onboarding path selected based on role and department",
            "Predicted 94% automation success rate",
            "Estimated 2.3 days reduction in manual processing time"
        ]
    }

# Initialize chat if empty
if not st.session_state.messages:
    agent_message("👋 Hello! I'm ARIA, your AI HR assistant. I'm ready to process new hire documents with advanced intelligence. Upload a document to get started!", "info")

# Display chat
display_chat()

# Document processing workflow
if uploaded_file:
    status_placeholder.info("🧠 AI Agent Processing...")
    
    # Step 1: AI Document Analysis
    if st.session_state.agent_step == 0:
        extracted_text = ai_extract_text(uploaded_file, uploaded_file.name)
        
        if extracted_text:
            # Step 2: AI Field Extraction
            fields, confidence_scores = ai_analyze_and_extract(extracted_text)
            st.session_state.extracted_data = (fields, confidence_scores)
            st.session_state.agent_step = 1
            
            # Display results
            st.markdown("---")
            st.markdown("## 🤖 AI Extraction Results")
            
            col1, col2 = st.columns(2)
            
            field_items = list(fields.items())
            mid_point = len(field_items) // 2
            
            with col1:
                for k, v in field_items[:mid_point]:
                    if "Not Found" in v:
                        st.markdown(f"**{k}**: :red[{v}]")
                    else:
                        confidence = confidence_scores.get(k, 0)
                        color = "green" if confidence > 0.8 else "orange" if confidence > 0.6 else "red"
                        st.markdown(f"**{k}**: :{color}[{v}] *({confidence:.1%} confidence)*")
            
            with col2:
                for k, v in field_items[mid_point:]:
                    if "Not Found" in v:
                        st.markdown(f"**{k}**: :red[{v}]")
                    else:
                        confidence = confidence_scores.get(k, 0)
                        color = "green" if confidence > 0.8 else "orange" if confidence > 0.6 else "red"
                        st.markdown(f"**{k}**: :{color}[{v}] *({confidence:.1%} confidence)*")
    
    # Step 3: AI Validation
    if st.session_state.agent_step == 1 and st.session_state.extracted_data:
        fields, confidence_scores = st.session_state.extracted_data
        
        if st.button("🧠 Run AI Validation", type="primary"):
            validation_results, errors, warnings, suggestions = ai_validate_data(fields, confidence_scores)
            
            st.markdown("---")
            st.markdown("## 🔍 AI Validation Results")
            
            if errors:
                st.error("❌ **Critical Issues Detected:**")
                for error in errors:
                    st.write(f"• {error}")
            
            if warnings:
                st.warning("⚠️ **AI Recommendations:**")
                for warning in warnings:
                    st.write(f"• {warning}")
            
            if suggestions:
                st.info("💡 **AI Optimization Suggestions:**")
                for suggestion in suggestions:
                    st.write(f"• {suggestion}")
            
            if not errors:
                st.success("🎉 **AI Validation Complete - Ready for ERP Integration!**")
                st.session_state.validation_complete = True
                st.session_state.agent_step = 2
    
    # Step 4: AI ERP Integration
    if st.session_state.agent_step == 2 and st.session_state.validation_complete:
        if st.button("🚀 Execute AI ERP Integration", type="primary"):
            fields, _ = st.session_state.extracted_data
            result = ai_erp_integration(fields)
            
            st.markdown("---")
            st.markdown("## 🎉 AI Integration Success!")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.success(f"**Employee ID:** {result['employee_id']}")
                st.info(f"**System:** {result['system']}")
                st.info(f"**Status:** {result['status']}")
            
            with col2:
                st.info(f"**Timestamp:** {result['timestamp']}")
                st.metric("AI Confidence", "97.8%", "↑ 1.2%")
                st.metric("Processing Speed", "12.4s", "↓ 3.1s")
            
            st.markdown("### 🧠 AI Insights")
            for insight in result['ai_insights']:
                st.markdown(f"• {insight}")
            
            # Reset for next document
            if st.button("🔄 Process Another Document"):
                st.session_state.agent_step = 0
                st.session_state.extracted_data = None
                st.session_state.validation_complete = False
                st.session_state.messages = []
                st.experimental_rerun()

else:
    status_placeholder.success("🟢 Ready to Assist")

# Clear chat button
if st.button("🗑️ Clear AI Chat"):
    st.session_state.messages = []
    st.experimental_rerun()
